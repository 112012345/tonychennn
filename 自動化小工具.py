import os
import glob
from docx import Document
import openpyxl
import re
from openpyxl.utils import column_index_from_string
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

def extract_table_after_heading(doc_path, heading_keyword):
    doc = Document(doc_path)
    start_find = False  # 先初始化

    for block in doc.element.body.iter():
        if block.tag.endswith('p') and block.text and heading_keyword in block.text:
            start_find = True
        if block.tag.endswith('tbl') and start_find:
            for table in doc.tables:
                if table._element == block:
                    data = []
                    print("--- 表格內容預覽 ---")
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = '\n'.join([para.text for para in cell.paragraphs])
                            print(cell_text)
                            row_data.append(cell_text)
                        data.append(row_data)
                    print("--- 表格結束 ---")
                    return data
    return None

def get_first_heading_or_nonempty_paragraph(doc_path):
    doc = Document(doc_path)
    # 先找第一個 Heading 1
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading 1') and para.text.strip():
            return para.text.strip()
    # 如果沒有 Heading 1，就找第一個非空段落
    for para in doc.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return None

def write_to_excel(excel_path, data, start_row=2, start_col=1, lon_value=None, column_name=None, first_title=None):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    # 1. 找所有目標欄位的最大行數
    target_cols = ['B', 'C', 'E', 'H', 'Z', 'AC', 'AF']
    max_row = 1
    for col in target_cols:
        col_idx = column_index_from_string(col)
        col_max = 1
        for r in range(ws.max_row, 0, -1):
            if ws.cell(row=r, column=col_idx).value not in (None, ""):
                col_max = r
                break
        if col_max > max_row:
            max_row = col_max

    # 2. 所有寫入 row_idx 的地方都從 max_row+1 開始
    row_idx_z = row_idx_ac = row_idx_af = row_idx_h = row_idx_b = row_idx_c = row_idx_e = max_row + 1

    # H 欄 append
    h_col = column_index_from_string('H')
    for i in range(1, len(data)):
        if len(data[i]) > 1:
            ws.cell(row=row_idx_h, column=h_col, value=data[i][1])
            row_idx_h += 1

    # B 欄 append LON-34, LON-34, ...
    if lon_value:
        b_col = column_index_from_string('B')
        for i in range(1, len(data)):
            ws.cell(row=row_idx_b, column=b_col, value=lon_value)
            row_idx_b += 1

    # C 欄 append first_title, first_title, ...
    if first_title:
        c_col = column_index_from_string('C')
        for i in range(1, len(data)):
            ws.cell(row=row_idx_c, column=c_col, value=first_title)
            row_idx_c += 1

    # Z/AC/AF 欄分流（每一格獨立判斷是否有方案二、方案三）
    z_col = column_index_from_string('Z')
    ac_col = column_index_from_string('AC')
    af_col = column_index_from_string('AF')
    split_token2 = "【方案二"
    split_token3 = "【方案三"
    for i in range(1, len(data)):
        if len(data[i]) > 2:
            cell_value = data[i][2]
            if split_token2 in cell_value and split_token3 in cell_value:
                idx2 = cell_value.index(split_token2)
                idx3 = cell_value.index(split_token3)
                before_plan2 = cell_value[:idx2]
                plan2_part = cell_value[idx2:idx3]
                plan3_part = cell_value[idx3:]
                # Z 欄寫前段
                ws.cell(row=row_idx_z, column=z_col, value=before_plan2)
                row_idx_z += 1
                # AC 欄寫方案二段
                ws.cell(row=row_idx_ac, column=ac_col, value=plan2_part)
                row_idx_ac += 1
                # AF 欄寫方案三段
                ws.cell(row=row_idx_af, column=af_col, value=plan3_part)
                row_idx_af += 1
            elif split_token2 in cell_value:
                idx2 = cell_value.index(split_token2)
                before_plan2 = cell_value[:idx2]
                after_plan2 = cell_value[idx2:]
                # Z 欄寫前段
                ws.cell(row=row_idx_z, column=z_col, value=before_plan2)
                row_idx_z += 1
                # AC 欄寫後段
                ws.cell(row=row_idx_ac, column=ac_col, value=after_plan2)
                row_idx_ac += 1
            elif split_token3 in cell_value:
                idx3 = cell_value.index(split_token3)
                before_plan3 = cell_value[:idx3]
                plan3_part = cell_value[idx3:]
                # Z 欄寫前段
                ws.cell(row=row_idx_z, column=z_col, value=before_plan3)
                row_idx_z += 1
                # AF 欄寫方案三段
                ws.cell(row=row_idx_af, column=af_col, value=plan3_part)
                row_idx_af += 1
            else:
                # 沒有方案二、三，全部寫到 Z
                ws.cell(row=row_idx_z, column=z_col, value=cell_value)
                row_idx_z += 1

    # E 欄 append LON-34-01, LON-34-02, ...
    if lon_value:
        e_col = column_index_from_string('E')
        for i in range(1, len(data)):
            suffix = f"{i:02d}"
            ws.cell(row=row_idx_e, column=e_col, value=f"{lon_value}-{suffix}")
            row_idx_e += 1

    # 寫入 LON-34 到指定欄位名稱下方（只寫一次）
    if lon_value and column_name and ws is not None:
        for cell in next(ws.iter_rows(min_row=1, max_row=1)):
            if cell.value == column_name:
                col_idx = cell.column
                if col_idx is None:
                    col_idx = 1
                elif not isinstance(col_idx, int):
                    try:
                        col_idx = int(col_idx)
                    except Exception:
                        col_idx = 1
                ws.cell(row=max_row + 1, column=col_idx, value=lon_value)
                break

    # === 新增：I欄加下拉選單 ===
    dropdown_options = [
        "法遵類", "客戶權益", "內規要求", "產品需求",
        "帳務類", "內部流程", "管理需求", "IT議題",
        "業務常規", "使用者介面"
    ]
    i_col_letter = 'I'
    start_row = 2  # 假設第1列是標題
    end_row = ws.max_row
    dv = DataValidation(type="list", formula1='"{}"'.format(",".join(dropdown_options)), allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"{i_col_letter}{start_row}:{i_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：J欄加下拉選單 ===
    j_dropdown_options = [
        "(5分)關鍵影響業務目標或法規要求",
        "(3分)重要但有替代方案或影響範圍較局部",
        "(2分)次要影響，與業務優化或內部效率相關",
        "(1分)影響極低，非必要但有加分效果"
    ]
    j_col_letter = 'J'
    j_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(j_dropdown_options)), allow_blank=True)
    ws.add_data_validation(j_dv)
    j_dv.add(f"{j_col_letter}{start_row}:{j_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：L欄加下拉選單 ===
    l_dropdown_options = [
        "(5分)高度影響系統或業務運作",
        "(3分)中度影響業務或系統",
        "(2分)輕微影響業務或系統",
        "(1分)幾乎無影響"
    ]
    l_col_letter = 'L'
    l_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(l_dropdown_options)), allow_blank=True)
    ws.add_data_validation(l_dv)
    l_dv.add(f"{l_col_letter}{start_row}:{l_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：N欄加下拉選單 ===
    n_dropdown_options = [
        "(5分)高風險（合規、財務、技術）",
        "(3分)中風險（可控但需關注）",
        "(2分)低風險（影響有限）",
        "(1分)極低風險（幾乎無影響）"
    ]
    n_col_letter = 'N'
    n_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(n_dropdown_options)), allow_blank=True)
    ws.add_data_validation(n_dv)
    n_dv.add(f"{n_col_letter}{start_row}:{n_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：T欄加下拉選單 ===
    t_dropdown_options = [
        "S", "M", "L", "XL",
        "無差異", "待評估", "需求取消", "SSR2", "SSR3", "合併差異"
    ]
    t_col_letter = 'T'
    t_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(t_dropdown_options)), allow_blank=True)
    ws.add_data_validation(t_dv)
    t_dv.add(f"{t_col_letter}{start_row}:{t_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：W欄加下拉選單 ===
    w_dropdown_options = [
        "Must Have : 必要的,不這樣做系統無法運作",
        "Should Have : 應該的,讓系統可以正常運作",
        "Could Have : 可以的,增加系統附加價值",
        "Would Like to have, but not this time : 額外的期望,但不是此時最需要的"
    ]
    w_col_letter = 'W'
    w_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(w_dropdown_options)), allow_blank=True)
    ws.add_data_validation(w_dv)
    w_dv.add(f"{w_col_letter}{start_row}:{w_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：W欄加註解 ===
    from openpyxl.comments import Comment
    w_comment_text = (
        "決議方案說明\n"
        "BA及SA基於業務及時程(資源)考量\n"
        "這個評估決議方案（基於MoSCoW原則）選出\n"
        "- Must Have :\n  必要的,不這樣做系統無法運作\n"
        "- Should Have :\n  應該的,讓系統可以正常運作\n"
        "- Could Have :\n  可以的,增加系統附加價值\n"
        "- Would Like to have, but not this time :\n  額外的期望,但不是此時最需要的"
    )
    for row in range(start_row, end_row + 1):
        cell = ws[f"{w_col_letter}{row}"]
        cell.comment = Comment(w_comment_text, "MoSCoW說明")
    # === 新增結束 ===

    # === 新增：X欄加下拉選單 ===
    x_dropdown_options = ["Y", "N"]
    x_col_letter = 'X'
    x_dv = DataValidation(type="list", formula1='"{}"'.format(",".join(x_dropdown_options)), allow_blank=True)
    ws.add_data_validation(x_dv)
    x_dv.add(f"{x_col_letter}{start_row}:{x_col_letter}{end_row}")
    # === 新增結束 ===

    # === 新增：X欄加註解 ===
    x_comment_text = (
        "決議方案說明\n這個需求架構是可以沿用的"
    )
    for row in range(start_row, end_row + 1):
        cell = ws[f"{x_col_letter}{row}"]
        cell.comment = Comment(x_comment_text, "說明")
    # === 新增結束 ===

    wb.save('output/03_SSR差異分析方案評分工作表_授信.xlsx')



def run_batch_process(doc_folder, excel_path, heading_keyword, column_name):
    # 取得所有 docx 檔案
    doc_paths = glob.glob(os.path.join(doc_folder, "*.docx"))
    results = []
    for doc_path in doc_paths:
        filename = os.path.basename(doc_path)
        match = re.search(r'(LON-\d+)', filename)
        lon_value = match.group(1) if match else None
        table_data = extract_table_after_heading(doc_path, heading_keyword)
        first_title = get_first_heading_or_nonempty_paragraph(doc_path)
        if table_data:
            write_to_excel(excel_path, table_data, lon_value=lon_value, column_name=column_name, first_title=first_title)
            results.append((filename, '成功'))
        else:
            results.append((filename, '未找到對應標題或表格'))
    return results

# =====================
# 主程式：批次處理多個 Word 檔案
# =====================
# doc_folder = '/Users/chenfutung/Desktop/intern/2025 summer intern/ibm國泰/SSR/SSR0/授信組'  # 你的word檔案資料夾
# excel_path = '/Users/chenfutung/Desktop/03_SSR差異分析方案評分工作表_授信.xlsx'
# heading_keyword = "差異及"
# column_name = "文件編號"  # 你要找的欄位名稱
#
# # 取得所有 docx 檔案
# doc_paths = glob.glob(os.path.join(doc_folder, "*.docx"))
#
# for doc_path in doc_paths:
#     filename = os.path.basename(doc_path)
#     match = re.search(r'(LON-\d+)', filename)
#     lon_value = match.group(1) if match else None
#     table_data = extract_table_after_heading(doc_path, heading_keyword)
#     first_title = get_first_heading_or_nonempty_paragraph(doc_path)
#     if table_data:
#         write_to_excel(excel_path, table_data, lon_value=lon_value, column_name=column_name, first_title=first_title)
#     else:
#         print(f"未找到對應標題或表格：{filename}")